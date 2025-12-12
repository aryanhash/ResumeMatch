"""
File handling utilities for resume parsing
"""
import io
from typing import Union
from pypdf import PdfReader  # SECURITY: Updated from PyPDF2
from docx import Document

# SECURITY: File size limits
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
ALLOWED_MIME_TYPES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'text/plain'
}


def extract_text_from_pdf(file_content: Union[bytes, io.BytesIO]) -> str:
    """Extract text from PDF file"""
    try:
        if isinstance(file_content, bytes):
            file_content = io.BytesIO(file_content)
        
        reader = PdfReader(file_content)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_content: Union[bytes, io.BytesIO]) -> str:
    """Extract text from DOCX file"""
    try:
        if isinstance(file_content, bytes):
            file_content = io.BytesIO(file_content)
        
        doc = Document(file_content)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def validate_file(file_content: bytes, filename: str, content_type: str = None) -> None:
    """
    Validate file before processing
    
    SECURITY: Validates file size, extension, and content type
    """
    # Check file size
    if len(file_content) > MAX_FILE_SIZE:
        raise ValueError(f"File too large. Maximum size: {MAX_FILE_SIZE / (1024*1024):.1f}MB")
    
    if len(file_content) == 0:
        raise ValueError("File is empty")
    
    # Check extension
    filename_lower = filename.lower()
    if not any(filename_lower.endswith(ext) for ext in ALLOWED_EXTENSIONS):
        raise ValueError(f"Unsupported file format. Allowed: {', '.join(ALLOWED_EXTENSIONS)}")
    
    # Check MIME type if provided
    if content_type and content_type not in ALLOWED_MIME_TYPES:
        raise ValueError(f"Unsupported MIME type: {content_type}")


def extract_text_from_file(file_content: bytes, filename: str, content_type: str = None) -> str:
    """
    Extract text from file based on extension
    
    SECURITY: Validates file before processing
    """
    # Validate file first
    validate_file(file_content, filename, content_type)
    
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_content)
    elif filename_lower.endswith('.txt'):
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            raise ValueError("File contains invalid UTF-8 characters")
    else:
        raise ValueError(f"Unsupported file format: {filename}")

