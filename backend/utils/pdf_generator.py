"""
PDF Resume Generator - Creates professional PDF resumes
"""
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from io import BytesIO
from typing import Dict, Any, List
import logging

logger = logging.getLogger(__name__)


class ResumePDFGenerator:
    """Generates professional PDF resumes"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for the resume"""
        # Name style - large, bold, centered
        self.styles.add(ParagraphStyle(
            name='ResumeName',
            parent=self.styles['Heading1'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            textColor=colors.black
        ))
        
        # Contact info style
        self.styles.add(ParagraphStyle(
            name='ContactInfo',
            parent=self.styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER,
            spaceAfter=12,
            textColor=colors.HexColor('#333333')
        ))
        
        # Section header style
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=11,
            fontName='Helvetica-Bold',
            spaceBefore=12,
            spaceAfter=4,
            textColor=colors.black,
            borderPadding=0
        ))
        
        # Job title style
        self.styles.add(ParagraphStyle(
            name='JobTitle',
            parent=self.styles['Normal'],
            fontSize=10,
            fontName='Helvetica-Bold',
            spaceBefore=8,
            spaceAfter=2
        ))
        
        # Company/Date style
        self.styles.add(ParagraphStyle(
            name='CompanyDate',
            parent=self.styles['Normal'],
            fontSize=10,
            fontName='Helvetica-Oblique',
            spaceAfter=4,
            textColor=colors.HexColor('#444444')
        ))
        
        # Bullet point style
        self.styles.add(ParagraphStyle(
            name='BulletPoint',
            parent=self.styles['Normal'],
            fontSize=10,
            leftIndent=12,
            spaceBefore=2,
            spaceAfter=2,
            bulletIndent=0,
            textColor=colors.HexColor('#333333')
        ))
        
        # Normal text style
        self.styles.add(ParagraphStyle(
            name='ResumeText',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            textColor=colors.HexColor('#333333'),
            alignment=TA_JUSTIFY
        ))
        
        # Skills style
        self.styles.add(ParagraphStyle(
            name='SkillsText',
            parent=self.styles['Normal'],
            fontSize=10,
            leftIndent=12,
            spaceAfter=2
        ))

    def generate(self, result: Dict[str, Any]) -> BytesIO:
        """Generate PDF resume from the processing result"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.5*inch,
            leftMargin=0.5*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        
        story = []
        parsed_resume = result.get('parsed_resume', {})
        rewritten_resume = result.get('rewritten_resume', {})
        
        # === HEADER ===
        name = parsed_resume.get('name', 'CANDIDATE NAME')
        story.append(Paragraph(name.upper(), self.styles['ResumeName']))
        
        # Contact info
        contact_parts = []
        if parsed_resume.get('location'):
            contact_parts.append(parsed_resume['location'])
        if parsed_resume.get('phone'):
            contact_parts.append(f"Ph No: {parsed_resume['phone']}")
        if parsed_resume.get('email'):
            contact_parts.append(parsed_resume['email'])
        
        if contact_parts:
            contact_line = " | ".join(contact_parts)
            story.append(Paragraph(contact_line, self.styles['ContactInfo']))
        
        # LinkedIn/GitHub
        links = []
        if parsed_resume.get('linkedin'):
            links.append(f"LinkedIn: {parsed_resume['linkedin']}")
        if parsed_resume.get('github'):
            links.append(f"GitHub: {parsed_resume['github']}")
        if links:
            story.append(Paragraph(" | ".join(links), self.styles['ContactInfo']))
        
        # === PROFESSIONAL SUMMARY ===
        if rewritten_resume.get('summary'):
            story.append(self._section_header("PROFESSIONAL SUMMARY"))
            story.append(Paragraph(rewritten_resume['summary'], self.styles['ResumeText']))
        
        # === EDUCATION ===
        education = parsed_resume.get('education', [])
        if education:
            story.append(self._section_header("EDUCATION"))
            for edu in education:
                degree_text = edu.get('degree', '')
                if edu.get('field_of_study'):
                    degree_text += f" in {edu['field_of_study']}"
                
                institution = edu.get('institution', '')
                location = edu.get('location', '')
                year = edu.get('year', '')
                gpa = edu.get('gpa', '')
                
                # Create two-column layout for education
                left_text = f"<b>{institution}</b><br/>{degree_text}"
                if gpa:
                    left_text += f"<br/>GPA/CGPA: {gpa}"
                
                right_text = f"{location}<br/>{year}" if location else year
                
                story.append(Paragraph(f"<b>{institution}</b>", self.styles['JobTitle']))
                story.append(Paragraph(f"{degree_text}", self.styles['CompanyDate']))
                if gpa:
                    story.append(Paragraph(f"GPA/CGPA: {gpa}", self.styles['BulletPoint']))
        
        # === WORK EXPERIENCE ===
        experience = parsed_resume.get('experience', [])
        if experience:
            story.append(self._section_header("WORK EXPERIENCE"))
            
            # Build improved bullets map
            improved_map = {}
            for item in rewritten_resume.get('improved_bullets', []):
                improved_map[item.get('original', '')] = item.get('improved', '')
            
            for exp in experience:
                title = exp.get('title', '')
                company = exp.get('company', '')
                duration = exp.get('duration', '')
                location = exp.get('location', '')
                
                story.append(Paragraph(f"<b>{company}</b>", self.styles['JobTitle']))
                
                subtitle_parts = [title]
                if location:
                    subtitle_parts.append(location)
                subtitle_parts.append(duration)
                story.append(Paragraph(" | ".join(filter(None, subtitle_parts)), self.styles['CompanyDate']))
                
                # Bullet points (use improved if available)
                for bullet in exp.get('description', []):
                    improved_bullet = improved_map.get(bullet, bullet)
                    story.append(Paragraph(f"• {improved_bullet}", self.styles['BulletPoint']))
        
        # === PROJECTS ===
        projects = parsed_resume.get('projects', [])
        if projects:
            story.append(self._section_header("PERSONAL PROJECTS"))
            for proj in projects:
                proj_name = proj.get('name', '')
                tech = proj.get('technologies', [])
                tech_str = f" – {', '.join(tech[:5])}" if tech else ""
                
                story.append(Paragraph(f"<b>{proj_name}{tech_str}</b>", self.styles['JobTitle']))
                
                if proj.get('description'):
                    story.append(Paragraph(f"• {proj['description']}", self.styles['BulletPoint']))
                if proj.get('impact'):
                    story.append(Paragraph(f"• Impact: {proj['impact']}", self.styles['BulletPoint']))
        
        # === TECHNICAL SKILLS ===
        skills = rewritten_resume.get('reordered_skills', parsed_resume.get('skills', []))
        if skills:
            story.append(self._section_header("TECHNICAL SKILLS"))
            
            # Group skills by category if possible, otherwise list them
            skills_text = ", ".join(skills[:20])  # Limit to top 20 skills
            story.append(Paragraph(f"• {skills_text}", self.styles['SkillsText']))
        
        # === CERTIFICATIONS ===
        certs = parsed_resume.get('certifications', [])
        if certs:
            story.append(self._section_header("CERTIFICATIONS / AWARDS"))
            for cert in certs:
                story.append(Paragraph(f"• {cert}", self.styles['BulletPoint']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _section_header(self, title: str):
        """Create a section header with underline"""
        elements = []
        elements.append(Spacer(1, 8))
        elements.append(Paragraph(title, self.styles['SectionHeader']))
        elements.append(HRFlowable(
            width="100%",
            thickness=1,
            color=colors.black,
            spaceBefore=2,
            spaceAfter=6
        ))
        return elements
    
    def _section_header(self, title: str):
        """Create a section header - returns paragraph with line"""
        from reportlab.platypus import KeepTogether
        return Paragraph(
            f'<u><b>{title}</b></u>',
            self.styles['SectionHeader']
        )


class ResumeWordGenerator:
    """Generates Word documents for resumes"""
    
    def generate(self, result: Dict[str, Any]) -> BytesIO:
        """Generate Word document from the processing result"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = Document()
        
        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        parsed_resume = result.get('parsed_resume', {})
        rewritten_resume = result.get('rewritten_resume', {})
        
        # === HEADER ===
        name = parsed_resume.get('name', 'CANDIDATE NAME')
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name.upper())
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if parsed_resume.get('location'):
            contact_parts.append(parsed_resume['location'])
        if parsed_resume.get('phone'):
            contact_parts.append(f"Ph No: {parsed_resume['phone']}")
        if parsed_resume.get('email'):
            contact_parts.append(parsed_resume['email'])
        
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(10)
        
        # === PROFESSIONAL SUMMARY ===
        if rewritten_resume.get('summary'):
            self._add_section_header(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(rewritten_resume['summary'])
            for run in summary_para.runs:
                run.font.size = Pt(10)
        
        # === EDUCATION ===
        education = parsed_resume.get('education', [])
        if education:
            self._add_section_header(doc, "EDUCATION")
            for edu in education:
                inst_para = doc.add_paragraph()
                inst_run = inst_para.add_run(edu.get('institution', ''))
                inst_run.bold = True
                inst_run.font.size = Pt(10)
                
                degree_text = edu.get('degree', '')
                if edu.get('field_of_study'):
                    degree_text += f" in {edu['field_of_study']}"
                
                degree_para = doc.add_paragraph(degree_text)
                for run in degree_para.runs:
                    run.font.size = Pt(10)
                    run.italic = True
        
        # === WORK EXPERIENCE ===
        experience = parsed_resume.get('experience', [])
        if experience:
            self._add_section_header(doc, "WORK EXPERIENCE")
            
            improved_map = {}
            for item in rewritten_resume.get('improved_bullets', []):
                improved_map[item.get('original', '')] = item.get('improved', '')
            
            for exp in experience:
                # Company name
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(exp.get('company', ''))
                company_run.bold = True
                company_run.font.size = Pt(10)
                
                # Title and duration
                title_parts = [exp.get('title', ''), exp.get('duration', '')]
                title_para = doc.add_paragraph(" | ".join(filter(None, title_parts)))
                for run in title_para.runs:
                    run.font.size = Pt(10)
                    run.italic = True
                
                # Bullets
                for bullet in exp.get('description', []):
                    improved = improved_map.get(bullet, bullet)
                    bullet_para = doc.add_paragraph(f"• {improved}")
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
                    for run in bullet_para.runs:
                        run.font.size = Pt(10)
        
        # === PROJECTS ===
        projects = parsed_resume.get('projects', [])
        if projects:
            self._add_section_header(doc, "PERSONAL PROJECTS")
            for proj in projects:
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(proj.get('name', ''))
                proj_run.bold = True
                proj_run.font.size = Pt(10)
                
                if proj.get('technologies'):
                    tech_run = proj_para.add_run(f" – {', '.join(proj['technologies'][:5])}")
                    tech_run.font.size = Pt(10)
                
                if proj.get('description'):
                    desc_para = doc.add_paragraph(f"• {proj['description']}")
                    desc_para.paragraph_format.left_indent = Inches(0.25)
                    for run in desc_para.runs:
                        run.font.size = Pt(10)
        
        # === TECHNICAL SKILLS ===
        skills = rewritten_resume.get('reordered_skills', parsed_resume.get('skills', []))
        if skills:
            self._add_section_header(doc, "TECHNICAL SKILLS")
            skills_para = doc.add_paragraph(f"• {', '.join(skills[:20])}")
            skills_para.paragraph_format.left_indent = Inches(0.25)
            for run in skills_para.runs:
                run.font.size = Pt(10)
        
        # === CERTIFICATIONS ===
        certs = parsed_resume.get('certifications', [])
        if certs:
            self._add_section_header(doc, "CERTIFICATIONS / AWARDS")
            for cert in certs:
                cert_para = doc.add_paragraph(f"• {cert}")
                cert_para.paragraph_format.left_indent = Inches(0.25)
                for run in cert_para.runs:
                    run.font.size = Pt(10)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _add_section_header(self, doc, title: str):
        """Add a section header with underline"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Add some space
        doc.add_paragraph()
        
        # Add header
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(title)
        header_run.bold = True
        header_run.font.size = Pt(11)
        
        # Add bottom border (underline effect)
        pPr = header_para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)


# Standalone functions for backward compatibility with main.py imports

def generate_resume_pdf(parsed_resume, rewritten_resume) -> bytes:
    """
    Generate optimized resume PDF from parsed resume and rewritten resume data.
    Uses LaTeX (Jake's Resume template) for professional one-page formatting.
    Falls back to ReportLab if LaTeX is not available.
    Returns PDF as bytes.
    """
    # Convert Pydantic models to dicts if needed
    result = {
        'parsed_resume': parsed_resume.model_dump() if hasattr(parsed_resume, 'model_dump') else parsed_resume,
        'rewritten_resume': rewritten_resume.model_dump() if hasattr(rewritten_resume, 'model_dump') else rewritten_resume
    }
    
    # Try LaTeX generator first (produces better formatting matching Jake's Resume template)
    try:
        from .latex_pdf_generator import LaTeXResumeGenerator
        generator = LaTeXResumeGenerator()
        buffer = generator.generate(result)
        return buffer.getvalue()
    except (RuntimeError, FileNotFoundError, ImportError) as e:
        logger.warning(f"LaTeX generator unavailable ({str(e)}), falling back to ReportLab")
        # Fall back to ReportLab
        generator = ResumePDFGenerator()
        buffer = generator.generate(result)
        return buffer.getvalue()


def generate_cover_letter_pdf(cover_letter_content: str, candidate_name: str = "Candidate") -> bytes:
    """
    Generate cover letter PDF.
    Returns PDF as bytes.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        name='CoverLetterTitle',
        parent=styles['Heading1'],
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=20,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        name='CoverLetterBody',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceBefore=6,
        spaceAfter=6,
        leading=14
    )
    
    story = []
    
    # Title
    story.append(Paragraph("COVER LETTER", title_style))
    story.append(Spacer(1, 20))
    
    # Split content into paragraphs and add them
    paragraphs = cover_letter_content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Handle single newlines as line breaks
            para = para.replace('\n', '<br/>')
            story.append(Paragraph(para, body_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_resume_word(parsed_resume, rewritten_resume) -> bytes:
    """
    Generate optimized resume as Word document.
    Returns DOCX as bytes.
    """
    result = {
        'parsed_resume': parsed_resume.model_dump() if hasattr(parsed_resume, 'model_dump') else parsed_resume,
        'rewritten_resume': rewritten_resume.model_dump() if hasattr(rewritten_resume, 'model_dump') else rewritten_resume
    }
    
    generator = ResumeWordGenerator()
    buffer = generator.generate(result)
    return buffer.getvalue()
